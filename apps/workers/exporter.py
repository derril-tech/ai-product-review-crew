# Created automatically by Cursor AI (2024-12-19)

import logging
from typing import Dict, List, Optional, Any
from datetime import datetime
import json
import io
from celery import shared_task
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import boto3
from botocore.exceptions import ClientError

logger = logging.getLogger(__name__)


@shared_task(bind=True, name='exporter.export_pdf')
def export_review_pdf(
    self,
    review_id: str,
    review_data: Dict[str, Any],
    export_config: Dict[str, Any]
) -> Dict[str, Any]:
    """
    Export review to PDF format.
    """
    try:
        logger.info("Starting PDF export", review_id=review_id)
        
        # Create PDF document
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center
        )
        title = Paragraph(review_data.get('title', 'Product Review'), title_style)
        story.append(title)
        story.append(Spacer(1, 20))
        
        # Executive Summary
        if 'executive_summary' in review_data.get('sections', {}):
            story.append(Paragraph("Executive Summary", styles['Heading2']))
            story.append(Spacer(1, 12))
            summary_text = review_data['sections']['executive_summary'].get('content', '')
            story.append(Paragraph(summary_text, styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Product Comparison Table
        if 'products' in review_data:
            story.append(Paragraph("Product Comparison", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            products = review_data['products']
            if products:
                # Create table data
                table_data = [['Product', 'Score', 'Price', 'Rank']]
                for product in products:
                    table_data.append([
                        product.get('name', ''),
                        f"{product.get('score', 0):.2f}",
                        f"${product.get('price', 0)}",
                        f"#{product.get('rank', 0)}"
                    ])
                
                # Create table
                table = Table(table_data, colWidths=[3*inch, 1*inch, 1*inch, 0.8*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
                story.append(Spacer(1, 20))
        
        # Detailed Analysis
        if 'detailed_analysis' in review_data.get('sections', {}):
            story.append(Paragraph("Detailed Analysis", styles['Heading2']))
            story.append(Spacer(1, 12))
            analysis_text = review_data['sections']['detailed_analysis'].get('content', '')
            story.append(Paragraph(analysis_text, styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Pros and Cons
        if 'pros' in review_data or 'cons' in review_data:
            story.append(Paragraph("Pros and Cons", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            if 'pros' in review_data:
                story.append(Paragraph("Pros:", styles['Heading3']))
                for pro in review_data['pros']:
                    story.append(Paragraph(f"• {pro.get('title', '')}", styles['Normal']))
                story.append(Spacer(1, 12))
            
            if 'cons' in review_data:
                story.append(Paragraph("Cons:", styles['Heading3']))
                for con in review_data['cons']:
                    story.append(Paragraph(f"• {con.get('title', '')}", styles['Normal']))
                story.append(Spacer(1, 20))
        
        # Conclusion
        if 'conclusion' in review_data.get('sections', {}):
            story.append(Paragraph("Conclusion", styles['Heading2']))
            story.append(Spacer(1, 12))
            conclusion_text = review_data['sections']['conclusion'].get('content', '')
            story.append(Paragraph(conclusion_text, styles['Normal']))
        
        # Build PDF
        doc.build(story)
        pdf_content = buffer.getvalue()
        buffer.close()
        
        # Upload to S3/MinIO
        file_key = f"exports/{review_id}/review_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.pdf"
        upload_success = _upload_to_storage(pdf_content, file_key, 'application/pdf')
        
        logger.info("PDF export completed", review_id=review_id, file_key=file_key)
        
        return {
            'review_id': review_id,
            'format': 'pdf',
            'file_key': file_key,
            'file_size': len(pdf_content),
            'upload_success': upload_success,
            'metadata': {
                'exported_at': datetime.utcnow().isoformat(),
                'pages': len(story) // 2  # Rough estimate
            }
        }
        
    except Exception as e:
        logger.error("Error in PDF export", review_id=review_id, error=str(e))
        raise self.retry(countdown=60, max_retries=3)


@shared_task(bind=True, name='exporter.export_word')
def export_review_word(
    self,
    review_id: str,
    review_data: Dict[str, Any],
    export_config: Dict[str, Any]
) -> Dict[str, Any]:
    """
    Export review to Word document format.
    """
    try:
        logger.info("Starting Word export", review_id=review_id)
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading(review_data.get('title', 'Product Review'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive Summary
        if 'executive_summary' in review_data.get('sections', {}):
            doc.add_heading('Executive Summary', level=1)
            summary_text = review_data['sections']['executive_summary'].get('content', '')
            doc.add_paragraph(summary_text)
        
        # Product Comparison Table
        if 'products' in review_data:
            doc.add_heading('Product Comparison', level=1)
            
            products = review_data['products']
            if products:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                # Header row
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Product'
                header_cells[1].text = 'Score'
                header_cells[2].text = 'Price'
                header_cells[3].text = 'Rank'
                
                # Data rows
                for product in products:
                    row_cells = table.add_row().cells
                    row_cells[0].text = product.get('name', '')
                    row_cells[1].text = f"{product.get('score', 0):.2f}"
                    row_cells[2].text = f"${product.get('price', 0)}"
                    row_cells[3].text = f"#{product.get('rank', 0)}"
        
        # Detailed Analysis
        if 'detailed_analysis' in review_data.get('sections', {}):
            doc.add_heading('Detailed Analysis', level=1)
            analysis_text = review_data['sections']['detailed_analysis'].get('content', '')
            doc.add_paragraph(analysis_text)
        
        # Pros and Cons
        if 'pros' in review_data or 'cons' in review_data:
            doc.add_heading('Pros and Cons', level=1)
            
            if 'pros' in review_data:
                doc.add_heading('Pros:', level=2)
                for pro in review_data['pros']:
                    doc.add_paragraph(pro.get('title', ''), style='List Bullet')
            
            if 'cons' in review_data:
                doc.add_heading('Cons:', level=2)
                for con in review_data['cons']:
                    doc.add_paragraph(con.get('title', ''), style='List Bullet')
        
        # Conclusion
        if 'conclusion' in review_data.get('sections', {}):
            doc.add_heading('Conclusion', level=1)
            conclusion_text = review_data['sections']['conclusion'].get('content', '')
            doc.add_paragraph(conclusion_text)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_content = buffer.getvalue()
        buffer.close()
        
        # Upload to S3/MinIO
        file_key = f"exports/{review_id}/review_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
        upload_success = _upload_to_storage(docx_content, file_key, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        
        logger.info("Word export completed", review_id=review_id, file_key=file_key)
        
        return {
            'review_id': review_id,
            'format': 'docx',
            'file_key': file_key,
            'file_size': len(docx_content),
            'upload_success': upload_success,
            'metadata': {
                'exported_at': datetime.utcnow().isoformat(),
                'sections': len(doc.paragraphs)
            }
        }
        
    except Exception as e:
        logger.error("Error in Word export", review_id=review_id, error=str(e))
        raise self.retry(countdown=60, max_retries=3)


@shared_task(bind=True, name='exporter.export_html')
def export_review_html(
    self,
    review_id: str,
    review_data: Dict[str, Any],
    export_config: Dict[str, Any]
) -> Dict[str, Any]:
    """
    Export review to HTML format.
    """
    try:
        logger.info("Starting HTML export", review_id=review_id)
        
        # Generate HTML content
        html_content = _generate_html_content(review_data, export_config)
        
        # Upload to S3/MinIO
        file_key = f"exports/{review_id}/review_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.html"
        upload_success = _upload_to_storage(html_content.encode('utf-8'), file_key, 'text/html')
        
        logger.info("HTML export completed", review_id=review_id, file_key=file_key)
        
        return {
            'review_id': review_id,
            'format': 'html',
            'file_key': file_key,
            'file_size': len(html_content),
            'upload_success': upload_success,
            'metadata': {
                'exported_at': datetime.utcnow().isoformat(),
                'include_styles': export_config.get('include_styles', True)
            }
        }
        
    except Exception as e:
        logger.error("Error in HTML export", review_id=review_id, error=str(e))
        raise self.retry(countdown=60, max_retries=3)


@shared_task(bind=True, name='exporter.export_json')
def export_review_json(
    self,
    review_id: str,
    review_data: Dict[str, Any],
    export_config: Dict[str, Any]
) -> Dict[str, Any]:
    """
    Export review to JSON format.
    """
    try:
        logger.info("Starting JSON export", review_id=review_id)
        
        # Prepare JSON data
        json_data = {
            'review_id': review_id,
            'exported_at': datetime.utcnow().isoformat(),
            'review': review_data,
            'export_config': export_config
        }
        
        # Convert to JSON string
        json_content = json.dumps(json_data, indent=2, ensure_ascii=False)
        
        # Upload to S3/MinIO
        file_key = f"exports/{review_id}/review_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.json"
        upload_success = _upload_to_storage(json_content.encode('utf-8'), file_key, 'application/json')
        
        logger.info("JSON export completed", review_id=review_id, file_key=file_key)
        
        return {
            'review_id': review_id,
            'format': 'json',
            'file_key': file_key,
            'file_size': len(json_content),
            'upload_success': upload_success,
            'metadata': {
                'exported_at': datetime.utcnow().isoformat(),
                'data_keys': list(review_data.keys())
            }
        }
        
    except Exception as e:
        logger.error("Error in JSON export", review_id=review_id, error=str(e))
        raise self.retry(countdown=60, max_retries=3)


def _generate_html_content(review_data: Dict[str, Any], export_config: Dict[str, Any]) -> str:
    """Generate HTML content for the review."""
    include_styles = export_config.get('include_styles', True)
    
    html = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{review_data.get('title', 'Product Review')}</title>
        {_get_html_styles() if include_styles else ''}
    </head>
    <body>
        <div class="container">
            <h1 class="title">{review_data.get('title', 'Product Review')}</h1>
    """
    
    # Executive Summary
    if 'executive_summary' in review_data.get('sections', {}):
        html += f"""
            <section class="section">
                <h2>Executive Summary</h2>
                <p>{review_data['sections']['executive_summary'].get('content', '')}</p>
            </section>
        """
    
    # Product Comparison Table
    if 'products' in review_data:
        html += """
            <section class="section">
                <h2>Product Comparison</h2>
                <table class="comparison-table">
                    <thead>
                        <tr>
                            <th>Product</th>
                            <th>Score</th>
                            <th>Price</th>
                            <th>Rank</th>
                        </tr>
                    </thead>
                    <tbody>
        """
        
        for product in review_data['products']:
            html += f"""
                        <tr>
                            <td>{product.get('name', '')}</td>
                            <td>{product.get('score', 0):.2f}</td>
                            <td>${product.get('price', 0)}</td>
                            <td>#{product.get('rank', 0)}</td>
                        </tr>
            """
        
        html += """
                    </tbody>
                </table>
            </section>
        """
    
    # Detailed Analysis
    if 'detailed_analysis' in review_data.get('sections', {}):
        html += f"""
            <section class="section">
                <h2>Detailed Analysis</h2>
                <p>{review_data['sections']['detailed_analysis'].get('content', '')}</p>
            </section>
        """
    
    # Pros and Cons
    if 'pros' in review_data or 'cons' in review_data:
        html += """
            <section class="section">
                <h2>Pros and Cons</h2>
        """
        
        if 'pros' in review_data:
            html += """
                <div class="pros-cons">
                    <h3>Pros:</h3>
                    <ul>
            """
            for pro in review_data['pros']:
                html += f"<li>{pro.get('title', '')}</li>"
            html += "</ul></div>"
        
        if 'cons' in review_data:
            html += """
                <div class="pros-cons">
                    <h3>Cons:</h3>
                    <ul>
            """
            for con in review_data['cons']:
                html += f"<li>{con.get('title', '')}</li>"
            html += "</ul></div>"
        
        html += "</section>"
    
    # Conclusion
    if 'conclusion' in review_data.get('sections', {}):
        html += f"""
            <section class="section">
                <h2>Conclusion</h2>
                <p>{review_data['sections']['conclusion'].get('content', '')}</p>
            </section>
        """
    
    html += """
        </div>
    </body>
    </html>
    """
    
    return html


def _get_html_styles() -> str:
    """Get CSS styles for HTML export."""
    return """
    <style>
        body {
            font-family: Arial, sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
            background-color: #f5f5f5;
        }
        .container {
            max-width: 800px;
            margin: 0 auto;
            background-color: white;
            padding: 40px;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        .title {
            text-align: center;
            color: #333;
            border-bottom: 3px solid #007bff;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }
        .section {
            margin-bottom: 30px;
        }
        .section h2 {
            color: #007bff;
            border-left: 4px solid #007bff;
            padding-left: 15px;
        }
        .section h3 {
            color: #555;
            margin-top: 20px;
        }
        .comparison-table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }
        .comparison-table th,
        .comparison-table td {
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }
        .comparison-table th {
            background-color: #007bff;
            color: white;
            font-weight: bold;
        }
        .comparison-table tr:nth-child(even) {
            background-color: #f9f9f9;
        }
        .pros-cons {
            margin: 20px 0;
        }
        .pros-cons ul {
            margin: 10px 0;
            padding-left: 20px;
        }
        .pros-cons li {
            margin: 5px 0;
        }
    </style>
    """


def _upload_to_storage(content: bytes, file_key: str, content_type: str) -> bool:
    """Upload file to S3/MinIO storage."""
    try:
        # Configure S3 client (works with MinIO too)
        s3_client = boto3.client(
            's3',
            endpoint_url='http://localhost:9000',  # MinIO endpoint
            aws_access_key_id='minioadmin',
            aws_secret_access_key='minioadmin',
            region_name='us-east-1'
        )
        
        # Upload file
        s3_client.put_object(
            Bucket='product-review-crew',
            Key=file_key,
            Body=content,
            ContentType=content_type
        )
        
        return True
        
    except ClientError as e:
        logger.error(f"Error uploading to storage: {e}")
        return False
    except Exception as e:
        logger.error(f"Unexpected error uploading to storage: {e}")
        return False
